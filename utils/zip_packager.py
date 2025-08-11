import zipfile
import io
from typing import Dict
from datetime import datetime
from typing import Optional

try:
    from docx import Document  # type: ignore
    from docx.shared import Mm  # type: ignore
    from docx.enum.section import WD_ORIENT  # type: ignore
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup  # type: ignore
    BS4_AVAILABLE = True
except Exception:
    BS4_AVAILABLE = False

class ZipPackager:
    """Handles packaging of documents into ZIP files"""
    
    def create_package(self, documents: Dict[str, str], pdf_files: Dict[str, bytes], merged_pdf: bytes) -> io.BytesIO:
        """
        Create a ZIP package containing all documents in multiple formats
        
        Args:
            documents: Dictionary of HTML documents
            pdf_files: Dictionary of PDF files
            merged_pdf: Merged PDF content
            
        Returns:
            ZIP file as BytesIO buffer
        """
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Add HTML documents
            for doc_name, html_content in documents.items():
                filename = f"html/{doc_name.replace(' ', '_').lower()}.html"
                zip_file.writestr(filename, html_content)
            
            # Add individual PDF files
            for pdf_name, pdf_content in pdf_files.items():
                filename = f"pdf/{pdf_name}"
                zip_file.writestr(filename, pdf_content)
            
            # Add merged PDF
            zip_file.writestr("combined/all_documents_combined.pdf", merged_pdf)
            
            # Add Word documents generated from the same HTML
            for doc_name, html_content in documents.items():
                filename = f"word/{doc_name.replace(' ', '_').lower()}.docx"
                docx_bytes = self._html_to_docx_bytes(doc_name, html_content)
                zip_file.writestr(filename, docx_bytes)
        
        zip_buffer.seek(0)
        return zip_buffer

    def _html_to_docx_bytes(self, doc_name: str, html: str) -> bytes:
        """Convert HTML content to a DOCX file in-memory.
        Attempts a faithful representation: text, simple headings, lists, and tables.
        Ensures A4 with 10mm margins; landscape for Deviation Statement.
        """
        # If python-docx is not available, return the HTML bytes as a fallback placeholder
        if not DOCX_AVAILABLE:
            return html.encode("utf-8")

        document = Document()
        # Page setup: A4 with 10 mm margins. Landscape for Deviation Statement
        section = document.sections[0]
        is_landscape = ("deviation" in doc_name.lower())
        if is_landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Mm(297)
            section.page_height = Mm(210)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Mm(210)
            section.page_height = Mm(297)
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)

        if BS4_AVAILABLE:
            soup = BeautifulSoup(html, "html.parser")

            def is_within(node, ancestor_name):
                return node.find_parent(ancestor_name) is not None

            # Title/header if present
            header_div = soup.find(class_="header")
            if header_div:
                for node in header_div.find_all(recursive=False):
                    text = node.get_text(strip=True)
                    if not text:
                        continue
                    if node.name in ("h1", "h2", "h3", "h4", "h5", "h6") or "title" in node.get("class", []):
                        level = 1 if node.name == "h1" else 2 if node.name == "h2" else 3
                        document.add_heading(text, level=level)
                    elif "subtitle" in node.get("class", []):
                        document.add_paragraph(text)
                    else:
                        document.add_paragraph(text)

            # Tables (preserve order roughly by iterating through DOM)
            # We'll process elements in body order: handle tables inline, otherwise handle blocks
            body = soup.body or soup
            for elem in body.descendants:
                if not getattr(elem, 'name', None):
                    continue
                # Skip anything inside header; we already handled header
                if header_div and (elem is header_div or elem.find_parent(class_="header")):
                    continue

                if elem.name == "table":
                    # Determine number of columns from the first row
                    rows = elem.find_all("tr")
                    if not rows:
                        continue
                    first_cells = rows[0].find_all(["th", "td"]) or []
                    num_cols = max(1, len(first_cells))
                    docx_table = document.add_table(rows=0, cols=num_cols)
                    for tr in rows:
                        cells = tr.find_all(["th", "td"]) or []
                        row_cells = docx_table.add_row().cells
                        for idx in range(num_cols):
                            cell_text = cells[idx].get_text(strip=False) if idx < len(cells) else ""
                            row_cells[idx].text = " ".join(cell_text.split())
                    continue

                # Headings
                if elem.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                    if is_within(elem, "table"):
                        continue
                    text = elem.get_text(strip=True)
                    if text:
                        level_map = {"h1":1, "h2":2, "h3":3, "h4":4, "h5":5, "h6":6}
                        document.add_heading(text, level=level_map.get(elem.name, 2))
                    continue

                # Lists
                if elem.name in ("ul", "ol"):
                    if is_within(elem, "table"):
                        continue
                    for li in elem.find_all("li", recursive=False):
                        text = li.get_text(strip=True)
                        if text:
                            style = "List Bullet" if elem.name == "ul" else "List Number"
                            try:
                                document.add_paragraph(text, style=style)
                            except Exception:
                                document.add_paragraph(text)
                    continue

                # Paragraph-like blocks
                if elem.name in ("p", "div"):
                    if is_within(elem, "table") or is_within(elem, "ul") or is_within(elem, "ol"):
                        continue
                    text = elem.get_text(strip=True)
                    if text:
                        document.add_paragraph(text)
        else:
            # Best-effort fallback: strip basic HTML tags
            plain = html
            for tag in ["<br>", "<br/>", "<br />"]:
                plain = plain.replace(tag, "\n")
            # Remove other tags roughly
            import re as _re
            plain = _re.sub(r"<[^>]+>", "", plain)
            for line in [ln.strip() for ln in plain.splitlines() if ln.strip()]:
                document.add_paragraph(line)

        output = io.BytesIO()
        document.save(output)
        return output.getvalue()
